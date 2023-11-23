import os
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from PIL import Image
import pytesseract


def extract_text_from_image(image_path):
    try:
        with Image.open(image_path) as img:
            text = pytesseract.image_to_string(img, lang='eng')
            return text
    except Exception as e:
        print(f"Error processing {image_path}: {str(e)}")
        return None


def add_image_and_text(doc, image_path, text):
    doc.add_picture(image_path, width=Inches(5))

    if text is not None:
        paragraph = doc.add_paragraph(text)
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    doc.add_paragraph()


def generate_word_file(folder_path, output_file):
    doc = Document()

    for root, dirs, files in os.walk(folder_path):
        for file_name in files:
            if file_name.lower().endswith(('.png', '.jpg', '.jpeg', '.gif', '.bmp')):
                image_path = os.path.join(root, file_name)

                image_text = extract_text_from_image(image_path)

                add_image_and_text(doc, image_path, image_text)

    doc.save(output_file)


# Example usage
input_image_folder_path = 'folder_path'
output_word_file_path = 'output_file.docx'
generate_word_file(input_image_folder_path, output_word_file_path)
